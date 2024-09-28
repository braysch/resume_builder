# -*- coding: utf-8 -*-
from PyInquirer import prompt
from docx import Document
from datetime import datetime

# Global variables to store selected expertise and skills
selected_expertise = []
selected_skills = []

# Function to replace placeholders in paragraphs, even if split across runs
def replace_text_in_paragraphs(paragraphs, placeholders):
    for paragraph in paragraphs:
        full_text = ''.join([run.text for run in paragraph.runs])

        for key, value in placeholders.items():
            if key in full_text:
                new_text = full_text.replace(key, value)

                # Clear the original runs and update them with the new text
                for run in paragraph.runs:
                    run.text = ''  # Clear each run

                paragraph.runs[0].text = new_text


# Function to replace placeholders inside a table (even for nested tables)
def replace_text_in_tables(tables, placeholders):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(cell.paragraphs, placeholders)
                if cell.tables:
                    replace_text_in_tables(cell.tables, placeholders)


# Prompt for Position Title
position = input("Enter the position: ")
title = input("Enter your title (leave blank to use position): ")
if title == "":
    title = position
company = input("Enter the company name: ")
adjective = "innovative"
adjective = input("What adjective should we use for your resume summary?: ")
adjective.capitalize()

summary = "" + adjective + " " + title.lower() + " with professional experience in full-stack development for growing business. Highly knowledgeable in an array of languages, frameworks, and tools for web and mobile development to deliver helpful, effective, and visually pleasing applications."

# List of expertise areas and corresponding time strings
expertise = [
    ("Programming", "7y"),
    ("Mobile Development", "3y"),
    ("Web Development", "3y"),
    ("UI Design", "3y"),
    ("Back-end Development", "3y"),
    ("API Integration", "1y"),
    ("Embedded Systems", "4y"),
    ("Database Management", "2y"),
]

# List of skills and corresponding time strings
skills = [
    ("Android Studio", "3y"),
    ("Flutter", "2y"),
    ("Dart", "2y"),
    ("Java", "4y"),
    ("Kotlin", "3y"),
    ("Jetpack Compose", "3y"),
    ("React", "1y"),
    ("Vue", "2y"),
    ("JavaScript", "3y"),
    ("TypeScript", "1y"),
    ("HTML", "3y"),
    ("CSS", "3y"),
    ("Tailwind", "1y"),
    ("PHP", "1y"),
    ("Laravel", "1y"),
    ("C", "2y"),
    ("C#", "2y"),
    ("C++", "7y"),
    ("SQL", "2y"),
    ("Git", "4y"),
    ("Google Firebase", "3y"),
    ("Bootstrap", "1y"),
    ("Python", "2y"),
    ("Visual Studio", "5y"),
]

softSkills = [
    "Innovation",
    "Creativity",
    "Adaptability",
    "Resilience",
    "Problem Solving",
    "Critical Thinking"
]

# Function to prioritize selected expertise items
def prioritize_expertise():
    global selected_expertise  # Declare the global variable
    questions = [
        {
            'type': 'checkbox',
            'name': 'expertise',
            'message': 'Select and prioritize your expertise (use space to select, enter to confirm):',
            'choices': [{'name': exp[0]} for exp in expertise],
        },
    ]
    answers = prompt(questions)

    # Separate the selected items from the unselected ones
    selected_expertise = answers['expertise']
    unselected = [exp for exp in expertise if exp[0] not in selected_expertise]

    # Combine the selected and unselected lists
    return selected_expertise + [exp[0] for exp in unselected]


# Function to prioritize selected skills
def prioritize_skills():
    global selected_skills  # Declare the global variable
    questions = [
        {
            'type': 'checkbox',
            'name': 'skills',
            'message': 'Select and prioritize your skills (use space to select, enter to confirm):',
            'choices': [{'name': skill[0]} for skill in skills],
        },
    ]
    answers = prompt(questions)

    # Separate the selected items from the unselected ones
    selected_skills = answers['skills']
    unselected = [skill for skill in skills if skill[0] not in selected_skills]

    # Combine the selected and unselected lists
    return selected_skills + [skill[0] for skill in unselected]

# Function to gather soft skills from the user
def gather_soft_skills():
    # Prompt user for soft skills
    soft_skills_input = input("Enter your soft skills, separated by commas: ")
    soft_skills_list = [skill.strip() for skill in soft_skills_input.split(",") if skill.strip()]

    # Check if fewer than 6 soft skills are provided
    while len(soft_skills_list) < 6:
        print("You need at least 6 soft skills. Please select " + str(6-len(soft_skills_list)) + " more from the list.")
        questions = [
            {
                'type': 'checkbox',
                'name': 'additional_soft_skills',
                'message': 'Select additional soft skills (use space to select, enter to confirm):',
                'choices': [{'name': skill} for skill in softSkills if skill not in soft_skills_list],
            },
        ]
        answers = prompt(questions)
        soft_skills_list += answers['additional_soft_skills']
        # Remove duplicates and ensure a maximum of 6 soft skills
        soft_skills_list = list(dict.fromkeys(soft_skills_list))[:6]

    return soft_skills_list

# Load the template
template_path = 'resume_template.docx'
try:
    doc = Document(template_path)
    print(f"Template '{template_path}' loaded successfully.")
except Exception as e:
    print(f"Error loading template: {e}")
    exit()

# Get prioritized expertise and skills
prioritized_expertise = prioritize_expertise()
prioritized_skills = prioritize_skills()

# Gather soft skills
soft_skills = gather_soft_skills()

# Dictionary of placeholders and values
placeholders = {
    '{{Title}}': title,
    '{{Summary}}': summary,
    '{{Company}}': company,
}

# Add expertise placeholders and their times to the dictionary
for i, exp in enumerate(prioritized_expertise):
    # Get the corresponding expertise tuple from the original expertise list
    exp_tuple = next((e for e in expertise if e[0] == exp), None)

    if exp_tuple:  # Only proceed if the expertise tuple was found
        placeholders[f'{{{{expertise{i}}}}}'] = exp_tuple[0]  # Expertise name
        placeholders[f'{{{{ey{i}}}}}'] = exp_tuple[1]  # Time placeholder for expertise

# Add expertise placeholders and their times to the dictionary
for i, skill in enumerate(prioritized_skills):
    # Get the corresponding expertise tuple from the original expertise list
    skill_tuple = next((s for s in skills if s[0] == skill), None)

    if skill_tuple:  # Only proceed if the expertise tuple was found
        placeholders[f'{{{{skill{i}}}}}'] = skill_tuple[0]  # Expertise name
        placeholders[f'{{{{sy{i}}}}}'] = skill_tuple[1]  # Time placeholder for expertise

# Add soft skill placeholders to the dictionary
for i, soft_skill in enumerate(soft_skills):
    placeholders[f'{{{{softSkill{i}}}}}'] = soft_skill

# Replace placeholders in paragraphs outside tables
replace_text_in_paragraphs(doc.paragraphs, placeholders)

# Replace placeholders in tables (including nested tables)
replace_text_in_tables(doc.tables, placeholders)

# Save the updated document
output_path = 'resume.docx'
try:
    doc.save(output_path)
    print(f"Resume created and saved as '{output_path}'")
except Exception as e:
    print(f"Error saving document: {e}")

myQualities = input("...while using my [insert qualities]: ")
mission = input("...to further your mission of [insert mission]: ")
dontKnow = input("Job qualifications you don't have: ")

pythonStory = "I'm actually using Python right now to generate this cover letter content (but don’t worry. It’s still written by me and not a bot!) " if "Python" in selected_skills else ""
kotlinStory = "I originally began developing applications using Kotlin and Jetpack Compose, including an app to keep track of goals and routines." if "Kotlin" in selected_skills else ""
flutterStory = ("I have a published Flutter web application, a personal wedding website, which I’ve used to display event details, provide links to external sites, and gather guest and RSVP information in Firebase. " + ("One of my more ambitious Flutter projects is an app which allows a user to write code using a click-and-select user interface, convert the objective code into MicroPython, and use that code to program a microcontroller via wifi. " if "Java" not in selected_skills else "")) if ("Flutter" in selected_skills or "Dart" in selected_skills) else ""
reactStory = "I am currently working on finishing up a personal resume website using React, TypeScript, and Tailwind; all tools which I picked up about a year ago." if ("React" in selected_skills or "Tailwind" in selected_skills or "TypeScript" in selected_skills) else ""
phpStory = "At my current job, I was tasked with writing code to have my company’s website communicate with the APIs of two different cell providers. At the beginning of the project, I had little to no knowledge of PHP, Laravel, or APIs; but now that feature is fully implemented. " if ("PHP" in selected_skills or "Laravel" in selected_skills or "API Integration" in selected_expertise) else ""
javaStory = "For my senior project I developed an app in Java designed for novice programmers to use a click-and-select style UI to create and send code to a microcontroller board. I have since then rewritten the app from scratch with Flutter." if ("Java" in selected_skills) else ""

coverLetter = """I am writing to express my interest in fulfilling your vacant role of """ + position + """.\nI’m a """ + title + """ with 8+ year of programming experience and 4+ years of web/mobile dev experience. I'm in search of a position that strongly aligns with my passion and drive for designing innovative tech, and I believe this position would provide an opportunity for me to engage in fulfilling work while using my """ + myQualities + """ to further """ + company + """’s mission of """ + mission + """.\n
I dropped out of my first programming class in middle school, but years later, I was able to get back on the saddle via game development. I minored in Computer Science at Utah State where I learned fundamental programming languages such as C++, Java, and Python. """ + pythonStory + """During the last few years of my college experience, I took a special interest in web and mobile development. I cultivated an array of skills including HTML, JavaScript, CSS, Vue, Java, Kotlin, and Jetpack Compose.\n
I continued to maintain and expand my skill set as a developer after graduation, undertaking person web and mobile app projects using Android Studio. """ + kotlinStory + """After independently learning Flutter / Dart, I began to develop mobile applications, using Google Firebase for database management . """ + flutterStory + """ It’s important for me to write practical, neat, reusable code as well as provide an intuitive and aesthetic user experience.\n
In addition to my person experience, I also have professional experience building web applications using React, Bootstrap, Tailwind, and Vue. """ + reactStory + """I have professional experience with backend development using C, JavaScript, and TypeScript; database management using SQL, and API integration using PHP Laravel. """ + phpStory + """For all my projects, personal and professional, I have a strict history of managing my version control using Git.\n
Your job listing mentions """ + dontKnow + """, of which I have limited professional experience with; however, I have no doubt that I’ll be able to learn and apply these skills in the same way I’ve done with many others.\n
I would love to continue to discuss this position and how my expansive skill set would make me a great asset to your team. I would also be more than happy to demonstrate any software projects of mine to exemplify my ambition and competency. Please reach out to me for any questions concerning my candidacy. I look forward to hearing from you."""

# Load the template
template_path = 'cover_letter_template.docx'
try:
    doc = Document(template_path)
    print(f"Template '{template_path}' loaded successfully.")
except Exception as e:
    print(f"Error loading template: {e}")
    exit()

# Dictionary of placeholders and values for the cover letter
placeholders = {
    '{{Company}}': company,
    '{{Position}}': position,
    '{{Title}}': title,
    '{{CoverLetter}}': coverLetter,
    '{{Date}}': datetime.now().strftime("%B %d, %Y")
}

# Replace placeholders in paragraphs outside tables
replace_text_in_paragraphs(doc.paragraphs, placeholders)

# Replace placeholders in tables (including nested tables)
replace_text_in_tables(doc.tables, placeholders)

# Save the updated document
output_path = 'cover_letter.docx'
try:
    doc.save(output_path)
    print(f"Cover letter created and saved as '{output_path}'")
except Exception as e:
    print(f"Error saving document: {e}")